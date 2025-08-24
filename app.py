from flask import Flask, request, jsonify, render_template, send_file, Response, make_response
import openai
import os
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
import logging
import re
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# ---------------------------
# Load environment variables
# ---------------------------
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
print("OPENAI_API_KEY:", openai.api_key)

# ---------------------------
# Flask App & Global Variables
# ---------------------------
app = Flask(__name__)
last_answer = ""  # store last GPT answer globally
def extract_title(text):
    # Extract first line as title
    lines = text.strip().split("\n")
    if lines:
        title = lines[0]
        # remove unwanted chars for filename
        title = re.sub(r'[\\/*?:"<>|]',"",title)
        return title[:50] if len(title)>50 else title
    return "GPT_Output"
    
    
# ---------------------------
# Helper Functions
# ---------------------------
def extract_title(content: str) -> str:
    """
    Extracts a safe filename from GPT output.
    Uses <h1> if available, else first line, replaces non-alphanumeric characters.
    """
    try:
        soup = BeautifulSoup(content, "html.parser")
        h1 = soup.find("h1")
        if h1:
            return re.sub(r'\W+', '_', h1.text.strip())[:50] or "Generated_Content"
        return re.sub(r'\W+', '_', content.strip().split("\n")[0])[:50] or "Generated_Content"
    except Exception:
        return "Generated_Content"

# ---------------------------
# Routes
# ---------------------------
@app.route('/')
def index():
    return render_template("landing.html")

@app.route('/chatbot')
def chatbot():
    return render_template("chat.html")

@app.route('/docqa')
def docqa():
    return render_template("docqa.html")

@app.route('/comingsoon')
def comingsoon():
    return render_template("comingsoon.html")
# ---------------------------
# API To call Chat App
# ---------------------------
@app.route('/ask', methods=['POST'])
def ask():
    global last_answer
    question = request.form.get("question")
    if not question:
        return jsonify({"error": "No question provided"}), 400

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": question}]
        )
        answer = response.choices[0].message['content']
        last_answer = answer  # store globally
        return jsonify({"answer": answer})
    except Exception as e:
        logging.error(f"Error in /ask: {str(e)}")
        return jsonify({"error": str(e)}), 500
# ---------------------------
# Download As Doc
# ---------------------------
@app.route("/download-doc", methods=["POST"])
def download_doc():
    global last_answer
    try:
        doc = Document()
        lines = last_answer.splitlines()

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Main title (# Title)
            if line.startswith("# "):
                text = line.lstrip("# ").strip()
                para = doc.add_paragraph(text)
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(20)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Subheading (## Subheading)
            elif line.startswith("## "):
                text = line.lstrip("#").strip()
                para = doc.add_paragraph(text)
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(16)

            # Bullet points (- or * at start)
            elif line.startswith("- ") or line.startswith("* "):
                text = line[2:].strip()
                para = doc.add_paragraph(text, style='List Bullet')

            # Bold inline (**text**) anywhere in line
            else:
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = para.add_run(part)

        # Save document to memory
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = "Delhi_Article.docx"
        response = make_response(buf.getvalue())
        response.headers['Content-Disposition'] = f'attachment; filename={filename}'
        response.mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return response

    except Exception as e:
        logging.error(f"Error in download-doc: {str(e)}")
        return jsonify({"error": str(e)}), 500
# ---------------------------
# Download As PDF
# ---------------------------
@app.route("/download-pdf", methods=["POST"])
def download_pdf():
    global last_answer
    try:
        buf = BytesIO()
        pdf = canvas.Canvas(buf, pagesize=A4)
        width, height = A4
        pdf.setFont("Helvetica", 12)
        
        # Split content into lines for PDF
        lines = last_answer.splitlines()
        y = height - 50  # start from top

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Handle headings
            if line.startswith("# "):
                pdf.setFont("Helvetica-Bold", 20)
                pdf.drawString(50, y, line[2:].strip())
                y -= 30
                pdf.setFont("Helvetica", 12)
            elif line.startswith("## "):
                pdf.setFont("Helvetica-Bold", 16)
                pdf.drawString(50, y, line[3:].strip())
                y -= 25
                pdf.setFont("Helvetica", 12)
            elif line.startswith("- ") or line.startswith("* "):
                pdf.drawString(70, y, "• " + line[2:].strip())
                y -= 18
            else:
                # Inline bold (**text**) handled as plain text for simplicity
                clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                pdf.drawString(50, y, clean_line)
                y -= 18

            # Add new page if we reach bottom
            if y < 50:
                pdf.showPage()
                pdf.setFont("Helvetica", 12)
                y = height - 50

        pdf.save()
        buf.seek(0)

        response = make_response(buf.getvalue())
        response.headers['Content-Disposition'] = 'attachment; filename=Delhi_Article.pdf'
        response.mimetype = "application/pdf"
        return response

    except Exception as e:
        logging.error(f"Error in download-pdf: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route("/test_openai")
def test_openai():
    try:
        r = openai.Model.list()
        return f"OpenAI Models Count: {len(r['data'])}"
    except Exception as e:
        return str(e)
# ---------------------------
# Run App
# ---------------------------
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    app.run(host="0.0.0.0", port=5000, debug=True)
